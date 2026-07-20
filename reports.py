"""
📊 GENERADOR DE REPORTES - PDF y Excel
Reportes analíticos para toma de decisiones
"""

import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from typing import List, Dict, Optional
import streamlit as st
from fpdf import FPDF
import io
import json
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats

# ============================================================================
# GENERADOR DE REPORTES PDF
# ============================================================================

class PDFReportGenerator:
    """Genera reportes PDF profesionales para toma de decisiones"""

    def __init__(self):
        self.pdf = FPDF()
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.pdf.add_page()
        self.pdf.set_font("Arial", size=10)
        self.colors = {
            'header': (102, 126, 234),  # Púrpura
            'accent': (118, 75, 162),
            'text': (38, 39, 48),
            'light': (240, 242, 246)
        }

    def add_title(self, title: str, subtitle: str = ""):
        """Añade título al reporte"""
        self.pdf.set_font("Arial", "B", 20)
        self.pdf.set_text_color(*self.colors['header'])
        self.pdf.cell(0, 15, title, ln=True, align="C")

        if subtitle:
            self.pdf.set_font("Arial", "I", 12)
            self.pdf.set_text_color(*self.colors['accent'])
            self.pdf.cell(0, 10, subtitle, ln=True, align="C")

        self.pdf.set_text_color(*self.colors['text'])
        self.pdf.ln(5)

    def add_section_title(self, title: str):
        """Añade título de sección"""
        self.pdf.set_font("Arial", "B", 14)
        self.pdf.set_text_color(*self.colors['header'])
        self.pdf.cell(0, 10, title, ln=True)
        self.pdf.set_text_color(*self.colors['text'])
        self.pdf.ln(2)

    def add_metric(self, label: str, value: str, unit: str = ""):
        """Añade una métrica al reporte"""
        self.pdf.set_font("Arial", "", 11)
        self.pdf.cell(80, 8, label + ":", 0)

        self.pdf.set_font("Arial", "B", 11)
        self.pdf.cell(80, 8, f"{value} {unit}", ln=True)

    def add_table(self, data: List[List], headers: List[str], col_widths: List[int]):
        """Añade tabla al reporte"""
        self.pdf.set_font("Arial", "B", 10)
        self.pdf.set_fill_color(*self.colors['header'])
        self.pdf.set_text_color(255, 255, 255)

        for header, width in zip(headers, col_widths):
            self.pdf.cell(width, 10, header, border=1, fill=True, align="C")
        self.pdf.ln()

        self.pdf.set_font("Arial", "", 9)
        self.pdf.set_text_color(*self.colors['text'])

        for row in data:
            for cell, width in zip(row, col_widths):
                self.pdf.cell(width, 8, str(cell), border=1)
            self.pdf.ln()

        self.pdf.ln(3)

    def add_text(self, text: str):
        """Añade párrafo de texto"""
        self.pdf.set_font("Arial", "", 10)
        self.pdf.multi_cell(0, 5, text)
        self.pdf.ln(2)

    def add_timestamp(self):
        """Añade timestamp de generación"""
        self.pdf.set_font("Arial", "I", 8)
        self.pdf.set_text_color(150, 150, 150)
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        self.pdf.cell(0, 10, f"Generado: {timestamp}", align="R")

    def add_image(self, img_path_or_buf, w=150, h=0):
        """Añade una imagen al reporte centrada"""
        page_width = self.pdf.w
        x = (page_width - w) / 2
        
        if isinstance(img_path_or_buf, io.BytesIO):
            img = Image.open(img_path_or_buf)
        else:
            img = img_path_or_buf
            
        self.pdf.image(img, x=x, w=w, h=h)
        self.pdf.ln(5)

    def output_bytes(self) -> bytes:
        """Retorna el PDF como bytes"""
        return bytes(self.pdf.output())


def generate_consolidated_report(json_path: str = 'model_results.json') -> bytes:
    """Genera un reporte Word consolidado de todos los deportes leyendo el JSON de resultados."""
    try:
        with open(json_path, 'r') as f:
            all_results = json.load(f)
    except FileNotFoundError:
        return b""
        
    doc = Document()
    doc.add_heading('EVALUACIÓN CONSOLIDADA DE MODELOS (TIME-BASED SPLIT)', 0)
    
    p = doc.add_paragraph()
    p.add_run('Generado: ').bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    doc.add_paragraph('Este reporte contiene los resultados del entrenamiento y validación de los modelos de inteligencia artificial para los distintos deportes. La validación se realizó reservando estrictamente el último año de historial de partidos como conjunto de prueba.')
    
    for sport, result in all_results.items():
        doc.add_heading(f'Deporte: {sport.upper()}', level=1)
        best_model = result.get('best_model', 'N/A')
        
        p = doc.add_paragraph()
        p.add_run('Modelo Seleccionado (Mejor Rendimiento): ').bold = True
        p.add_run(best_model.upper())
        
        doc.add_heading('Comparativa de Modelos Candidatos', level=2)
        
        model_details = result.get('model_details', {})
        if model_details:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Modelo'
            hdr_cells[1].text = 'Accuracy'
            hdr_cells[2].text = 'F1 Score'
            hdr_cells[3].text = 'Tiempo (s)'
            
            for m_name, details in model_details.items():
                row_cells = table.add_row().cells
                prefix = "⭐ " if m_name == best_model else ""
                row_cells[0].text = f"{prefix}{m_name.upper()}"
                row_cells[1].text = f"{details.get('accuracy', 0):.4f}"
                row_cells[2].text = f"{details.get('f1', 0):.4f}"
                row_cells[3].text = f"{details.get('time_s', 0):.2f}"
                
            doc.add_paragraph()
            
            doc.add_heading('Hiperparámetros Óptimos', level=2)
            for m_name, details in model_details.items():
                p = doc.add_paragraph()
                p.add_run(f"{m_name.upper()}: ").bold = True
                params = details.get('params', {})
                if params:
                    p.add_run(str(params))
                else:
                    p.add_run("N/A o Default")
                    
    byte_io = io.BytesIO()
    doc.save(byte_io)
    return byte_io.getvalue()


class ExcelReportGenerator:
    """Genera reportes Excel con datos y formatos profesionales"""

    def __init__(self):
        self.wb = Workbook()
        self.ws = self.wb.active
        self.styles = {
            'header_fill': PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid"),
            'header_font': Font(bold=True, color="FFFFFF", size=12),
            'accent_fill': PatternFill(start_color="F0F2F6", end_color="F0F2F6", fill_type="solid"),
            'accent_font': Font(bold=True, size=11),
            'border': Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            ),
            'center_align': Alignment(horizontal='center', vertical='center'),
            'number_format': '0.00'
        }

    def write_dataframe(self, df: pd.DataFrame, sheet_name: str, start_row: int = 1):
        """Escribe un DataFrame en la hoja de Excel"""
        ws = self.wb[sheet_name] if sheet_name in self.wb.sheetnames else self.wb.create_sheet(sheet_name)

        # Escribir headers
        for col_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(row=start_row, column=col_idx)
            cell.value = col_name
            cell.font = self.styles['header_font']
            cell.fill = self.styles['header_fill']
            cell.alignment = self.styles['center_align']
            cell.border = self.styles['border']

        # Escribir datos
        for row_idx, row_data in enumerate(dataframe_to_rows(df, index=False, header=False), start_row + 1):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.value = value
                cell.border = self.styles['border']

                # Aplicar formato a números
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    cell.number_format = self.styles['number_format']

        # Ajustar ancho de columnas
        for col_idx, col_name in enumerate(df.columns, 1):
            ws.column_dimensions[chr(64 + col_idx)].width = max(len(str(col_name)) + 2, 15)

    def add_metric_sheet(self, metrics: Dict[str, any], sheet_name: str = "Métricas"):
        """Añade hoja con métricas principales"""
        if sheet_name not in self.wb.sheetnames:
            ws = self.wb.create_sheet(sheet_name, 0)
        else:
            ws = self.wb[sheet_name]

        ws.title = sheet_name
        row = 1

        for label, value in metrics.items():
            cell_label = ws.cell(row=row, column=1)
            cell_label.value = label
            cell_label.font = self.styles['accent_font']
            cell_label.fill = self.styles['accent_fill']

            cell_value = ws.cell(row=row, column=2)
            cell_value.value = value
            cell_value.font = Font(bold=True, size=11)

            row += 1

        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20

    def output_bytes(self) -> bytes:
        """Retorna el Excel como bytes"""
        output = io.BytesIO()
        self.wb.save(output)
        output.seek(0)
        return output.getvalue()


class WordReportGenerator:
    """Genera reportes en formato Word (.docx)"""

    def __init__(self):
        self.doc = Document()
        
    def add_title(self, title: str, subtitle: str = ""):
        p = self.doc.add_heading(title, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if subtitle:
            p2 = self.doc.add_paragraph(subtitle)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section_title(self, title: str):
        self.doc.add_heading(title, level=1)

    def add_metric(self, label: str, value: str, unit: str = ""):
        p = self.doc.add_paragraph()
        p.add_run(f"{label}: ")
        p.add_run(f"{value} {unit}").bold = True

    def add_table(self, data: List[List], headers: List[str]):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
        for row in data:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
        self.doc.add_paragraph()

    def add_text(self, text: str):
        self.doc.add_paragraph(text)

    def add_timestamp(self):
        p = self.doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def output_bytes(self) -> bytes:
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()


# ============================================================================
# REPORTES ESPECÍFICOS
# ============================================================================

def generate_predictions_report(predictions: List[Dict], report_type: str = 'pdf') -> bytes:
    """
    Genera reporte de predicciones

    Args:
        predictions: Lista de predicciones
        report_type: 'pdf' o 'excel'

    Returns:
        Bytes del archivo generado
    """
    df = pd.DataFrame(predictions)

    if report_type == 'pdf':
        pdf = PDFReportGenerator()
        pdf.add_title("REPORTE DE PREDICCIONES", "Análisis Detallado")

        # Métricas principales
        pdf.add_section_title("Métricas Principales")
        total_preds = len(df)
        correct_preds = len(df[df['prediction_status'] == 'won']) if 'prediction_status' in df.columns else 0
        accuracy = (correct_preds / total_preds * 100) if total_preds > 0 else 0
        avg_confidence = df['confidence_level'].mean() if 'confidence_level' in df.columns else 0

        pdf.add_metric("Total de Predicciones", str(total_preds))
        pdf.add_metric("Predicciones Correctas", str(correct_preds))
        pdf.add_metric("Tasa de Precisión", f"{accuracy:.2f}", "%")
        pdf.add_metric("Confianza Promedio", f"{avg_confidence:.2f}")

        # Tabla de predicciones
        pdf.add_section_title("Historial de Predicciones")
        if not df.empty:
            data = []
            cols_to_show = ['match_id', 'predicted_home_score', 'predicted_away_score', 'confidence_level']
            cols_available = [col for col in cols_to_show if col in df.columns]

            for _, row in df.head(10).iterrows():
                data.append([row.get(col, 'N/A') for col in cols_available])

            pdf.add_table(data, cols_available, [40, 30, 30, 30])

        # Análisis
        pdf.add_section_title("Análisis")
        pdf.add_text(
            f"Sobre un total de {total_preds} predicciones realizadas, se han acertado {correct_preds} "
            f"({accuracy:.1f}%). El nivel de confianza promedio es de {avg_confidence:.2f}. "
            f"Esto indica {'un buen desempeño' if accuracy > 60 else 'una oportunidad de mejora'}."
        )

        pdf.add_timestamp()
        return pdf.output_bytes()

    elif report_type == 'word':
        word = WordReportGenerator()
        word.add_title("REPORTE DE PREDICCIONES", "Análisis Detallado")

        word.add_section_title("Métricas Principales")
        total_preds = len(df)
        correct_preds = len(df[df['prediction_status'] == 'won']) if 'prediction_status' in df.columns else 0
        accuracy = (correct_preds / total_preds * 100) if total_preds > 0 else 0
        avg_confidence = df['confidence_level'].mean() if 'confidence_level' in df.columns else 0

        word.add_metric("Total de Predicciones", str(total_preds))
        word.add_metric("Predicciones Correctas", str(correct_preds))
        word.add_metric("Tasa de Precisión", f"{accuracy:.2f}", "%")
        word.add_metric("Confianza Promedio", f"{avg_confidence:.2f}")

        word.add_section_title("Historial de Predicciones")
        if not df.empty:
            data = []
            cols_to_show = ['match_id', 'predicted_home_score', 'predicted_away_score', 'confidence_level']
            cols_available = [col for col in cols_to_show if col in df.columns]

            for _, row in df.head(10).iterrows():
                data.append([row.get(col, 'N/A') for col in cols_available])

            word.add_table(data, cols_available)

        word.add_section_title("Análisis")
        word.add_text(
            f"Sobre un total de {total_preds} predicciones realizadas, se han acertado {correct_preds} "
            f"({accuracy:.1f}%). El nivel de confianza promedio es de {avg_confidence:.2f}. "
            f"Esto indica {'un buen desempeño' if accuracy > 60 else 'una oportunidad de mejora'}."
        )

        word.add_timestamp()
        return word.output_bytes()

    else:  # Excel
        excel = ExcelReportGenerator()

        # Métrica general
        metrics = {
            'Total de Predicciones': len(df),
            'Correctas': len(df[df['prediction_status'] == 'won']) if 'prediction_status' in df.columns else 0,
            'Tasa de Precisión': f"{(len(df[df['prediction_status'] == 'won']) / len(df) * 100):.2f}%" if len(df) > 0 else "0%",
            'Confianza Promedia': f"{df['confidence_level'].mean():.2f}" if 'confidence_level' in df.columns else "0.00"
        }

        excel.add_metric_sheet(metrics, "Resumen")
        excel.write_dataframe(df, "Predicciones Detalladas")

        return excel.output_bytes()


# Helper functions for generating matplotlib charts
def _generate_distribution_chart(data: List[float]) -> io.BytesIO:
    plt.figure(figsize=(6, 3))
    sns.histplot(data, bins=15, kde=False, color='#667EEA', stat='count', alpha=0.7)
    if len(data) > 1:
        mu, sigma = np.mean(data), np.std(data)
        if sigma > 0:
            x = np.linspace(min(data), max(data), 100)
            y = stats.norm.pdf(x, mu, sigma)
            bin_width = (max(data) - min(data)) / 15
            y_scaled = y * len(data) * bin_width
            plt.plot(x, y_scaled, color='#E53E3E', linewidth=2, label='Curva Normal')
    plt.title("Distribución de Confianza", fontsize=10, fontweight='bold', color='#262730')
    plt.xlabel("Confianza", fontsize=8)
    plt.ylabel("Frecuencia", fontsize=8)
    plt.grid(True, linestyle='--', alpha=0.5)
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close()
    buf.seek(0)
    return buf

def _generate_cdf_chart(data: List[float]) -> io.BytesIO:
    plt.figure(figsize=(6, 3))
    sorted_data = np.sort(data)
    cumulative = np.arange(1, len(sorted_data) + 1) / len(sorted_data)
    plt.plot(sorted_data, cumulative, color='#764BA2', linewidth=2, label='CDF')
    plt.title("Función de Distribución Acumulada (CDF)", fontsize=10, fontweight='bold', color='#262730')
    plt.xlabel("Confianza", fontsize=8)
    plt.ylabel("Probabilidad", fontsize=8)
    plt.grid(True, linestyle='--', alpha=0.5)
    plt.ylim(-0.05, 1.05)
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close()
    buf.seek(0)
    return buf

def _generate_time_series_chart(dates: List[datetime], values: List[float]) -> io.BytesIO:
    plt.figure(figsize=(8, 3))
    df = pd.DataFrame({'date': dates, 'value': values}).sort_values('date')
    plt.plot(df['date'], df['value'], marker='o', color='#667EEA', linewidth=1.5, label='Confianza')
    if len(df) > 7:
        df['ma_7'] = df['value'].rolling(window=7).mean()
        plt.plot(df['date'], df['ma_7'], color='#E53E3E', linestyle='--', linewidth=1.5, label='Media Móvil (7d)')
    plt.title("Evolución de Confianza en el Tiempo", fontsize=10, fontweight='bold', color='#262730')
    plt.xlabel("Fecha", fontsize=8)
    plt.ylabel("Nivel de Confianza", fontsize=8)
    plt.grid(True, linestyle='--', alpha=0.5)
    if len(df) > 7:
        plt.legend(fontsize=8)
    plt.xticks(rotation=15)
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close()
    buf.seek(0)
    return buf

def _generate_status_chart(statuses: List[str]) -> io.BytesIO:
    plt.figure(figsize=(6, 3))
    df_status = pd.Series(statuses).value_counts()
    
    colors = []
    for s in df_status.index:
        if s == 'won':
            colors.append('#48BB78')
        elif s == 'lost':
            colors.append('#E53E3E')
        else:
            colors.append('#A0AEC0')
            
    sns.barplot(x=df_status.index, y=df_status.values, hue=df_status.index, palette=colors, legend=False)
    plt.title("Distribución por Status", fontsize=10, fontweight='bold', color='#262730')
    plt.xlabel("Status", fontsize=8)
    plt.ylabel("Cantidad", fontsize=8)
    plt.grid(True, linestyle='--', alpha=0.5)
    plt.tight_layout()
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close()
    buf.seek(0)
    return buf


def generate_statistics_report(user_stats: Dict, predictions: List[Dict], report_type: str = 'pdf') -> bytes:
    """
    Genera reporte de estadísticas descriptivas

    Args:
        user_stats: Estadísticas del usuario
        predictions: Listado de predicciones
        report_type: 'pdf' o 'excel'

    Returns:
        Bytes del archivo generado
    """
    df_preds = pd.DataFrame(predictions)

    if report_type == 'pdf':
        pdf = PDFReportGenerator()
        pdf.add_title("REPORTE DE ESTADÍSTICAS DESCRIPTIVAS", "Análisis Completo del Desempeño")

        # Resumen Ejecutivo
        pdf.add_section_title("Resumen Ejecutivo")
        pdf.add_metric("Total de Predicciones", str(user_stats.get('total_predictions', 0)))
        pdf.add_metric("Predicciones Correctas", str(user_stats.get('correct_predictions', 0)))
        pdf.add_metric("Tasa de Precisión", f"{user_stats.get('accuracy_rate', 0):.2f}", "%")
        pdf.add_metric("Confianza Promedio", f"{user_stats.get('avg_confidence', 0):.2f}")
        pdf.add_metric("Ranking Global", str(user_stats.get('rank', 'N/A')))

        # Estadísticas Descriptivas
        pdf.add_section_title("Estadísticas Descriptivas")

        if not df_preds.empty and 'confidence_level' in df_preds.columns:
            confidence_values = df_preds['confidence_level'].dropna()

            pdf.add_metric("Confianza Máxima", f"{confidence_values.max():.2f}")
            pdf.add_metric("Confianza Mínima", f"{confidence_values.min():.2f}")
            pdf.add_metric("Desviación Estándar", f"{confidence_values.std():.2f}")
            pdf.add_metric("Mediana", f"{confidence_values.median():.2f}")

        # Análisis de Rendimiento
        pdf.add_section_title("Análisis de Rendimiento")

        if user_stats.get('accuracy_rate', 0) > 70:
            performance_text = "EXCELENTE: Tu precisión es superior al 70%. Continúa con tu estrategia actual."
        elif user_stats.get('accuracy_rate', 0) > 55:
            performance_text = "BUENO: Tu precisión está por encima del promedio. Hay oportunidad de mejora."
        else:
            performance_text = "DESARROLLO: Tu precisión está por debajo del 55%. Revisa tu estrategia de análisis."

        pdf.add_text(performance_text)

        # Recomendaciones
        pdf.add_section_title("Recomendaciones")
        recommendations = [
            "- Analiza tus predicciones incorrectas para identificar patrones",
            "- Incrementa confianza solo cuando tengas datos sólidos",
            "- Diversifica tus predicciones entre diferentes deportes",
            "- Mantén un registro detallado de tus análisis",
            "- Revisa regularmente tu rendimiento contra el benchmark del mercado"
        ]
        pdf.add_text("\n".join(recommendations))

        # --- SECCIÓN DE ANÁLISIS GRÁFICO (Nueva Página) ---
        if not df_preds.empty:
            pdf.pdf.add_page()
            pdf.add_section_title("Análisis Gráfico y Visualizaciones Avanzadas")

            # 1. Distribución de Confianza
            if 'confidence_level' in df_preds.columns:
                confidence_data = df_preds['confidence_level'].dropna().tolist()
                if confidence_data:
                    pdf.pdf.set_font("Arial", "B", 11)
                    pdf.pdf.cell(0, 8, "Distribución de Confianza (Histograma + Curva Normal)", ln=True)
                    dist_img = _generate_distribution_chart(confidence_data)
                    pdf.add_image(dist_img, w=130)

                    pdf.pdf.set_font("Arial", "B", 11)
                    pdf.pdf.cell(0, 8, "Función de Distribución Acumulada (CDF)", ln=True)
                    cdf_img = _generate_cdf_chart(confidence_data)
                    pdf.add_image(cdf_img, w=130)

            # 2. Serie Temporal y Status
            # Creamos una nueva página para los siguientes gráficos para no sobrecargar el layout
            pdf.pdf.add_page()
            pdf.add_section_title("Evolución Temporal y Análisis por Status")

            if 'created_at' in df_preds.columns:
                df_preds['created_at'] = pd.to_datetime(df_preds['created_at'], errors='coerce')
                df_sorted = df_preds.sort_values('created_at').dropna(subset=['created_at'])
                if not df_sorted.empty:
                    dates = df_sorted['created_at'].tolist()
                    values = df_sorted.get('confidence_level', pd.Series(range(len(df_sorted)))).tolist()
                    pdf.pdf.set_font("Arial", "B", 11)
                    pdf.pdf.cell(0, 8, "Evolución de Confianza en el Tiempo", ln=True)
                    time_img = _generate_time_series_chart(dates, values)
                    pdf.add_image(time_img, w=150)

            if 'prediction_status' in df_preds.columns:
                statuses = df_preds['prediction_status'].dropna().tolist()
                if statuses:
                    pdf.pdf.set_font("Arial", "B", 11)
                    pdf.pdf.cell(0, 8, "Distribución por Status de las Predicciones", ln=True)
                    status_img = _generate_status_chart(statuses)
                    pdf.add_image(status_img, w=130)

        pdf.add_timestamp()
        return pdf.output_bytes()

    elif report_type == 'word':
        word = WordReportGenerator()
        word.add_title("REPORTE DE ESTADÍSTICAS DESCRIPTIVAS", "Análisis Completo del Desempeño")

        word.add_section_title("Resumen Ejecutivo")
        word.add_metric("Total de Predicciones", str(user_stats.get('total_predictions', 0)))
        word.add_metric("Predicciones Correctas", str(user_stats.get('correct_predictions', 0)))
        word.add_metric("Tasa de Precisión", f"{user_stats.get('accuracy_rate', 0):.2f}", "%")
        word.add_metric("Confianza Promedio", f"{user_stats.get('avg_confidence', 0):.2f}")
        word.add_metric("Ranking Global", str(user_stats.get('rank', 'N/A')))

        word.add_section_title("Estadísticas Descriptivas")
        if not df_preds.empty and 'confidence_level' in df_preds.columns:
            confidence_values = df_preds['confidence_level'].dropna()
            word.add_metric("Confianza Máxima", f"{confidence_values.max():.2f}")
            word.add_metric("Confianza Mínima", f"{confidence_values.min():.2f}")
            word.add_metric("Desviación Estándar", f"{confidence_values.std():.2f}")
            word.add_metric("Mediana", f"{confidence_values.median():.2f}")

        word.add_section_title("Análisis de Rendimiento")
        if user_stats.get('accuracy_rate', 0) > 70:
            performance_text = "EXCELENTE: Tu precisión es superior al 70%. Continúa con tu estrategia actual."
        elif user_stats.get('accuracy_rate', 0) > 55:
            performance_text = "BUENO: Tu precisión está por encima del promedio. Hay oportunidad de mejora."
        else:
            performance_text = "DESARROLLO: Tu precisión está por debajo del 55%. Revisa tu estrategia de análisis."
        word.add_text(performance_text)

        word.add_section_title("Recomendaciones")
        recommendations = [
            "• Analiza tus predicciones incorrectas para identificar patrones",
            "• Incrementa confianza solo cuando tengas datos sólidos",
            "• Diversifica tus predicciones entre diferentes deportes",
            "• Mantén un registro detallado de tus análisis",
            "• Revisa regularmente tu rendimiento contra el benchmark del mercado"
        ]
        word.add_text("\n".join(recommendations))

        word.add_timestamp()
        return word.output_bytes()

    else:  # Excel
        excel = ExcelReportGenerator()

        # Hoja de métricas
        metrics = {
            'Total Predicciones': user_stats.get('total_predictions', 0),
            'Predicciones Correctas': user_stats.get('correct_predictions', 0),
            'Tasa de Precisión %': user_stats.get('accuracy_rate', 0),
            'Confianza Promedio': round(user_stats.get('avg_confidence', 0), 2),
            'Ranking Global': str(user_stats.get('rank', 'N/A')),
        }

        excel.add_metric_sheet(metrics, "Resumen Ejecutivo")

        # Hoja de predicciones detalladas
        if not df_preds.empty:
            excel.write_dataframe(df_preds, "Predicciones")

        return excel.output_bytes()


def generate_competitions_report(competitions: List[Dict], report_type: str = 'pdf') -> bytes:
    """Genera reporte de competencias"""
    df = pd.DataFrame(competitions)

    if report_type == 'pdf':
        pdf = PDFReportGenerator()
        pdf.add_title("REPORTE DE COMPETENCIAS", "Resumen de Torneos Activos")

        pdf.add_section_title("Competencias Activas")
        pdf.add_metric("Total de Competencias", str(len(df)))

        if not df.empty:
            data = []
            for _, comp in df.head(10).iterrows():
                data.append([
                    comp.get('name', 'N/A'),
                    str(comp.get('current_participants', 0)),
                    f"${comp.get('prize_pool', 0):.2f}"
                ])

            pdf.add_table(data, ["Competencia", "Participantes", "Premio Total"], [60, 35, 35])

        pdf.add_timestamp()
        return pdf.output_bytes()

    elif report_type == 'word':
        word = WordReportGenerator()
        word.add_title("REPORTE DE COMPETENCIAS", "Resumen de Torneos Activos")
        word.add_section_title("Competencias Activas")
        word.add_metric("Total de Competencias", str(len(df)))

        if not df.empty:
            data = []
            for _, comp in df.head(10).iterrows():
                data.append([
                    comp.get('name', 'N/A'),
                    str(comp.get('current_participants', 0)),
                    f"${comp.get('prize_pool', 0):.2f}"
                ])
            word.add_table(data, ["Competencia", "Participantes", "Premio Total"])

        word.add_timestamp()
        return word.output_bytes()

    else:  # Excel
        excel = ExcelReportGenerator()
        excel.write_dataframe(df, "Competencias")
        return excel.output_bytes()


def generate_comprehensive_report(user_data: Dict, predictions: List[Dict],
                                  competitions: List[Dict], report_type: str = 'pdf') -> bytes:
    """
    Genera reporte comprensivo con todos los datos

    Args:
        user_data: Datos del usuario
        predictions: Predicciones
        competitions: Competencias
        report_type: 'pdf' o 'excel'

    Returns:
        Bytes del archivo generado
    """

    if report_type == 'pdf':
        pdf = PDFReportGenerator()
        pdf.add_title("REPORTE COMPRENSIVO", "Análisis Integral del Desempeño")
        pdf.add_text(f"Generado para: {user_data.get('username', 'Usuario')}")
        pdf.add_text(f"Período: Últimos 30 días")

        # SECCIÓN 1: RESUMEN EJECUTIVO
        pdf.add_section_title("RESUMEN EJECUTIVO")

        accuracy = user_data.get('accuracy_rate', 0)
        confidence = user_data.get('avg_confidence', 0)

        summary_text = f"""
Total de Predicciones: {user_data.get('total_predictions', 0)}
Tasa de Precisión: {accuracy:.2f}%
Confianza Promedio: {confidence:.2f}
Predicciones Correctas: {user_data.get('correct_predictions', 0)}
Ranking Global: {user_data.get('rank', 'N/A')}
        """

        pdf.add_text(summary_text)

        # SECCIÓN 2: ANÁLISIS DETALLADO
        pdf.add_section_title("ANÁLISIS DETALLADO")

        df_preds = pd.DataFrame(predictions) if predictions else pd.DataFrame()

        if not df_preds.empty:
            pdf.add_text(f"Total de registros analizados: {len(df_preds)}")

            # Tabla muestra
            if len(df_preds) > 0:
                data = []
                for _, pred in df_preds.head(5).iterrows():
                    data.append([
                        str(pred.get('match_id', 'N/A'))[:10],
                        f"{pred.get('predicted_home_score', 0)}-{pred.get('predicted_away_score', 0)}",
                        f"{pred.get('confidence_level', 0):.2f}"
                    ])

                pdf.add_table(data, ["Match ID", "Predicción", "Confianza"], [50, 50, 40])

        # SECCIÓN 3: COMPETENCIAS
        if competitions:
            pdf.add_section_title("COMPETENCIAS")
            pdf.add_metric("Competencias Activas", str(len(competitions)))

        # SECCIÓN 4: RECOMENDACIONES
        pdf.add_section_title("RECOMENDACIONES PARA MEJORA")

        if accuracy > 70:
            recommendation = "Tu desempeño es excelente. Mantén tu estrategia actual."
        elif accuracy > 55:
            recommendation = "Buen desempeño. Considera aumentar datos de análisis."
        else:
            recommendation = "Revisa tu metodología de análisis para mejorar resultados."

        pdf.add_text(recommendation)

        pdf.add_timestamp()
        return pdf.output_bytes()

    elif report_type == 'word':
        word = WordReportGenerator()
        word.add_title("REPORTE COMPRENSIVO", "Análisis Integral del Desempeño")
        word.add_text(f"Generado para: {user_data.get('username', 'Usuario')}")
        word.add_text(f"Período: Últimos 30 días")

        word.add_section_title("RESUMEN EJECUTIVO")
        accuracy = user_data.get('accuracy_rate', 0)
        confidence = user_data.get('avg_confidence', 0)
        summary_text = f"Total de Predicciones: {user_data.get('total_predictions', 0)}\n" \
                       f"Tasa de Precisión: {accuracy:.2f}%\n" \
                       f"Confianza Promedio: {confidence:.2f}\n" \
                       f"Predicciones Correctas: {user_data.get('correct_predictions', 0)}\n" \
                       f"Ranking Global: {user_data.get('rank', 'N/A')}"
        word.add_text(summary_text)

        word.add_section_title("ANÁLISIS DETALLADO")
        df_preds = pd.DataFrame(predictions) if predictions else pd.DataFrame()
        if not df_preds.empty:
            word.add_text(f"Total de registros analizados: {len(df_preds)}")
            if len(df_preds) > 0:
                data = []
                for _, pred in df_preds.head(5).iterrows():
                    data.append([
                        str(pred.get('match_id', 'N/A'))[:10],
                        f"{pred.get('predicted_home_score', 0)}-{pred.get('predicted_away_score', 0)}",
                        f"{pred.get('confidence_level', 0):.2f}"
                    ])
                word.add_table(data, ["Match ID", "Predicción", "Confianza"])

        if competitions:
            word.add_section_title("COMPETENCIAS")
            word.add_metric("Competencias Activas", str(len(competitions)))

        word.add_section_title("RECOMENDACIONES PARA MEJORA")
        if accuracy > 70:
            recommendation = "Tu desempeño es excelente. Mantén tu estrategia actual."
        elif accuracy > 55:
            recommendation = "Buen desempeño. Considera aumentar datos de análisis."
        else:
            recommendation = "Revisa tu metodología de análisis para mejorar resultados."
        word.add_text(recommendation)

        word.add_timestamp()
        return word.output_bytes()

    else:  # Excel
        excel = ExcelReportGenerator()

        # Métricas principales
        metrics = {
            'Total Predicciones': user_data.get('total_predictions', 0),
            'Tasa de Precisión %': round(user_data.get('accuracy_rate', 0), 2),
            'Confianza Promedio': round(user_data.get('avg_confidence', 0), 2),
            'Predicciones Correctas': user_data.get('correct_predictions', 0),
            'Ranking': str(user_data.get('rank', 'N/A'))
        }

        excel.add_metric_sheet(metrics)

        # Datos detallados
        if predictions:
            df_preds = pd.DataFrame(predictions)
            excel.write_dataframe(df_preds, "Predicciones Detalladas")

        if competitions:
            df_comps = pd.DataFrame(competitions)
            excel.write_dataframe(df_comps, "Competencias")

        return excel.output_bytes()

def generate_model_evaluation_report(eval_results: Dict, stats_results: Dict, report_type: str = 'pdf') -> bytes:
    """Genera reporte de evaluación y comparación de modelos"""
    
    records = []
    for model_name, res in eval_results.items():
        records.append({
            'Modelo': model_name,
            'Accuracy': f"{res.get('accuracy', 0):.4f}",
            'F1 Score': f"{res.get('f1', 0):.4f}",
            'Precision': f"{res.get('precision', 0):.4f}",
            'Recall': f"{res.get('recall', 0):.4f}",
            'Tiempo (s)': f"{res.get('time', 0):.2f}",
            'AUC': f"{res.get('roc_auc', 0):.4f}" if res.get('roc_auc') else "N/A"
        })
    df_eval = pd.DataFrame(records)
    best_model = stats_results.get('best_model', 'N/A')
    
    stats_text = [f"Mejor modelo en validación cruzada: {best_model}"]
    for k, v in stats_results.items():
        if k == 'best_model' or k == best_model:
            continue
        sig = "Significativo" if v['is_significant'] else "No Significativo"
        stats_text.append(f"vs {k}: p-value = {v['p_value']:.4f} ({sig})")
    
    if report_type == 'pdf':
        pdf = PDFReportGenerator()
        pdf.add_title("EVALUACIÓN DE MODELOS PREDICTIVOS", "Reporte Comparativo")
        
        pdf.add_section_title("Rendimiento de Modelos")
        headers = ["Modelo", "Accuracy", "F1 Score", "Precision", "Tiempo (s)"]
        data = [[row['Modelo'], row['Accuracy'], row['F1 Score'], row['Precision'], row['Tiempo (s)']] for _, row in df_eval.iterrows()]
        pdf.add_table(data, headers, [40, 30, 30, 30, 30])
        
        pdf.add_section_title("Significación Estadística")
        for line in stats_text:
            pdf.add_text(line)
            
        pdf.add_text("\nNota: P-value < 0.05 indica que la diferencia es estadísticamente significativa.")
        pdf.add_timestamp()
        return pdf.output_bytes()
        
    elif report_type == 'word':
        word = WordReportGenerator()
        word.add_title("EVALUACIÓN DE MODELOS PREDICTIVOS", "Reporte Comparativo")
        
        word.add_section_title("Rendimiento de Modelos")
        headers = ["Modelo", "Accuracy", "F1 Score", "Precision", "Recall", "Tiempo (s)", "AUC"]
        data = [[row[h] for h in headers] for _, row in df_eval.iterrows()]
        word.add_table(data, headers)
        
        word.add_section_title("Significación Estadística")
        for line in stats_text:
            word.add_text(line)
            
        word.add_text("\nNota: P-value < 0.05 indica que la diferencia es estadísticamente significativa.")
        word.add_timestamp()
        return word.output_bytes()
        
    else:  # Excel
        excel = ExcelReportGenerator()
        
        metrics = {'Mejor Modelo CV': best_model}
        for i, line in enumerate(stats_text[1:]):
            metrics[f'Prueba T {i+1}'] = line
            
        excel.add_metric_sheet(metrics, "Estadísticas")
        excel.write_dataframe(df_eval, "Comparativa")
        
        return excel.output_bytes()


# ============================================================================
# INTERFAZ DE STREAMLIT PARA REPORTES
# ============================================================================

def show_reports_ui(supabase_client=None):
    """Interfaz Streamlit para generar reportes"""

    st.title("📊 Generador de Reportes")

    report_type = st.radio("Tipo de Reporte", ["Predicciones", "Estadísticas Descriptivas", "Competencias", "Comprensivo"])
    file_format = st.radio("Formato", ["PDF", "Excel", "Word"])

    # Datos de ejemplo
    sample_predictions = [
        {'match_id': '001', 'predicted_home_score': 2, 'predicted_away_score': 1, 'confidence_level': 0.85, 'prediction_status': 'won'},
        {'match_id': '002', 'predicted_home_score': 1, 'predicted_away_score': 1, 'confidence_level': 0.65, 'prediction_status': 'lost'},
    ]

    sample_stats = {
        'total_predictions': 50,
        'correct_predictions': 35,
        'accuracy_rate': 70.0,
        'avg_confidence': 0.72,
        'rank': 'Top 15%'
    }

    sample_competitions = [
        {'name': 'Liga Predictor', 'current_participants': 245, 'prize_pool': 1000},
    ]

    if st.button("📥 Generar Reporte"):
        file_format_lower = file_format.lower()

        if report_type == "Predicciones":
            report_bytes = generate_predictions_report(sample_predictions, file_format_lower)
            ext = 'pdf' if file_format == 'PDF' else ('xlsx' if file_format == 'Excel' else 'docx')
            filename = f"predicciones_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"

        elif report_type == "Estadísticas Descriptivas":
            report_bytes = generate_statistics_report(sample_stats, sample_predictions, file_format_lower)
            ext = 'pdf' if file_format == 'PDF' else ('xlsx' if file_format == 'Excel' else 'docx')
            filename = f"estadisticas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"

        elif report_type == "Competencias":
            report_bytes = generate_competitions_report(sample_competitions, file_format_lower)
            ext = 'pdf' if file_format == 'PDF' else ('xlsx' if file_format == 'Excel' else 'docx')
            filename = f"competencias_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"

        else:  # Comprensivo
            report_bytes = generate_comprehensive_report(sample_stats, sample_predictions, sample_competitions, file_format_lower)
            ext = 'pdf' if file_format == 'PDF' else ('xlsx' if file_format == 'Excel' else 'docx')
            filename = f"reporte_comprensivo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"

        mime_type = "application/pdf"
        if file_format == 'Excel':
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif file_format == 'Word':
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        st.download_button(
            label=f"📥 Descargar {file_format}",
            data=report_bytes,
            file_name=filename,
            mime=mime_type
        )

        st.success(f"✅ Reporte generado exitosamente")


if __name__ == "__main__":
    show_reports_ui()
